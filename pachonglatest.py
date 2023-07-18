import concurrent.futures
import requests
from bs4 import BeautifulSoup
import re
from docx import Document
import threading

# 要爬取的网址
url = "https://cnhope.en.alibaba.com/productgrouplist-807068565/Gaming_monitor.html?spm=a2700.shop_plgr.88.69"

# 创建一个新的 Word 文档
doc = Document()

# 创建一个锁对象
lock = threading.Lock()

# 创建一个集合来存储已经处理过的链接
processed_links = set()

# 处理单个链接的函数
def process_link(link):
    href = link.get("href")
    if href.startswith("//"):
        href = "https:" + href  # 完善 URL
        lock.acquire()  # 获取锁
        try:
            if href in processed_links:  # 如果链接已经被处理过，那么跳过这个链接
                return
            # 将链接添加到已处理链接的集合中
            processed_links.add(href)
        finally:
            lock.release()  # 无论是否发生错误，最后都要释放锁

        link_response = requests.get(href)
        link_html_content = link_response.text
        link_soup = BeautifulSoup(link_html_content, "html.parser")
        title_tag = link_soup.find("title")
        if title_tag is not None:
            title = title_tag.text.strip()
            match = re.search(r' - Buy (.+?) on Alibaba.com', title)  # 提取关键词部分
            if match:
                keywords = match.group(1).split(',')
                title = title.replace(match.group(), '')  # 提取 "- Buy" 之前的信息
                output_str = "### " + title.strip() + "\n"  # 构建输出字符串
                output_str += "\n".join(keyword.strip() for keyword in keywords) + "\n\n"  # 构建关键词部分
                print(output_str)  # 在控制台输出
                # 获取锁，将结果添加到 Word 文档中，并保存文档
                lock.acquire()
                try:
                    doc.add_heading(title, level=3)
                    for keyword in keywords:
                        doc.add_paragraph(keyword)
                    doc.save("output.docx")
                finally:
                    lock.release()  # 无论是否发生错误，最后都要释放锁

# 发送 HTTP GET 请求获取网页内容
response = requests.get(url)
html_content = response.text

# 使用 BeautifulSoup 解析 HTML 内容
soup = BeautifulSoup(html_content, "html.parser")

# 寻找特定的元素或数据
# 在这个例子中，找到所有的链接，并筛选出以特定前缀开头的链接
links = soup.find_all("a", href=re.compile("//www.alibaba.com/product-detail"))

# 使用 ThreadPoolExecutor 并发处理链接，最多同时处理 4 个链接
with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
    future_to_output = {executor.submit(process_link, link): link for link in links}
    for future in concurrent.futures.as_completed(future_to_output):
        pass  # 不需要再这里处理结果，因为结果已经在 process_link 函数中处理了
